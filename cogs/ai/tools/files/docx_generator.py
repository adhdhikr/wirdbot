"""
Word Document Generator
Creates .docx files with support for LaTeX equation rendering.
"""
import asyncio
import logging
import os
import re
from io import BytesIO
from typing import Optional

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Word doc creation will be unavailable.")

try:
    from math2docx import add_math
    MATH2DOCX_AVAILABLE = True
except ImportError:
    MATH2DOCX_AVAILABLE = False
    logger.warning("math2docx not installed. LaTeX equations will be plain text.")
DISPLAY_MATH_PATTERN = re.compile(r'\$\$(.*?)\$\$', re.DOTALL)
INLINE_MATH_PATTERN = re.compile(r'\$([^\$\n]+?)\$')


async def create_word_doc(
    content: str,
    output_path: str,
    title: str = None,
    convert_latex: bool = True
) -> str:
    """
    Create a Word document from text content.
    
    Args:
        content: The text content to put in the document
        output_path: Where to save the .docx file
        title: Optional document title
        convert_latex: Whether to convert $...$ and $$...$$ to equations
    
    Returns:
        Path to created file or error message
    """
    if not DOCX_AVAILABLE:
        return "Error: python-docx is not installed. Cannot create Word documents."
    
    try:
        loop = asyncio.get_running_loop()
        result = await loop.run_in_executor(
            None, _create_doc_sync, content, output_path, title, convert_latex
        )
        return result
    except Exception as e:
        logger.error(f"Failed to create Word doc: {e}")
        return f"Error creating Word document: {e}"


def _create_doc_sync(
    content: str,
    output_path: str,
    title: str = None,
    convert_latex: bool = True
) -> str:
    """Synchronous document creation."""
    doc = Document()
    if title:
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if convert_latex and MATH2DOCX_AVAILABLE:
        _add_content_with_latex(doc, content)
    else:
        for para_text in content.split('\n\n'):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
BOLD_PATTERN = re.compile(r'\*\*(.+?)\*\*')
ITALIC_PATTERN = re.compile(r'\*(.+?)\*')

def _add_content_with_latex(doc: Document, content: str):
    """
    Add content to document, converting LaTeX equations using math2docx.
    Handles basic Markdown-style lists, headings, and mixed inline/display math.
    """
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('#'):
            level = len(line.split(' ')[0])
            if level <= 6:
                text = line[level:].strip()
                doc.add_heading(text, level=level)
                continue
        if line.startswith(('* ', '- ')):
            para = doc.add_paragraph(style='List Bullet')
            line = line[2:].strip()
        else:
            para = doc.add_paragraph()
        parts = DISPLAY_MATH_PATTERN.split(line)
        
        for i, part in enumerate(parts):
            if i % 2 == 1:
                _add_math_run(para, part.strip())
            else:
                if not part:
                    continue
                inline_parts = INLINE_MATH_PATTERN.split(part)
                for j, subpart in enumerate(inline_parts):
                    if j % 2 == 1:
                        _add_math_run(para, subpart.strip())
                    else:
                        if subpart:
                            _add_formatted_text(para, subpart)


def _add_formatted_text(para, text: str):
    """
    Parses text for Markdown bold (**) and italic (*).
    Adds runs to the paragraph with appropriate formatting.
    """
    bold_parts = BOLD_PATTERN.split(text)
    
    for k, bold_part in enumerate(bold_parts):
        if k % 2 == 1:
            run = para.add_run(bold_part)
            run.bold = True
        else:
            if not bold_part:
                continue
            italic_parts = ITALIC_PATTERN.split(bold_part)
            for m, italic_part in enumerate(italic_parts):
                if m % 2 == 1:
                    run = para.add_run(italic_part)
                    run.italic = True
                else:
                    if italic_part:
                        para.add_run(italic_part)


def _add_math_run(para, latex_str: str):
    """Helper to safely add a math run to a paragraph."""
    try:
        safe_latex = latex_str.replace(r'\vec', r'\mathbf')
        
        add_math(para, safe_latex)
    except Exception as e:
        logger.warning(f"Failed to render equation '{latex_str}': {e}")
        run = para.add_run(f'[{latex_str}]')
        run.italic = True
        run.font.name = 'Cambria Math'





async def create_word_doc_bytes(
    content: str,
    title: str = None,
    convert_latex: bool = True
) -> Optional[BytesIO]:
    """
    Create a Word document and return as bytes (for Discord upload).
    """
    if not DOCX_AVAILABLE:
        return None
    
    try:
        loop = asyncio.get_running_loop()
        result = await loop.run_in_executor(
            None, _create_doc_bytes_sync, content, title, convert_latex
        )
        return result
    except Exception as e:
        logger.error(f"Failed to create Word doc bytes: {e}")
        return None


def _create_doc_bytes_sync(
    content: str,
    title: str = None,
    convert_latex: bool = True
) -> BytesIO:
    """Create document and return as BytesIO."""
    doc = Document()
    
    if title:
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if convert_latex and MATH2DOCX_AVAILABLE:
        _add_content_with_latex(doc, content)
    else:
        for para_text in content.split('\n\n'):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
